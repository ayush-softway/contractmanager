from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

MD_FILE = "/Users/shivansh/Downloads/contract mgmt/Softway_MSA_Template.md"
OUT_FILE = "/Users/shivansh/Downloads/contract mgmt/Softway_MSA_Template.docx"
LOGO_FILE = "/Users/shivansh/Downloads/contract mgmt/logo.png"

doc = Document()

def set_columns(section, num_cols):
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), '720') # 0.5 inch space between columns

def setup_section(section):
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

    footer = section.footer
    if len(footer.paragraphs) == 0:
        footer_para = footer.add_paragraph()
    else:
        footer_para = footer.paragraphs[0]
        footer_para.clear()
        
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Softway Solutions Inc. dba Culture+ | 1801 Main St. Houston, TX 77002 | Governing Law: State of Texas | Page ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# Initialize first section (1 column for intro)
setup_section(doc.sections[0])
set_columns(doc.sections[0], 1)

# Default style
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(6)

def styled_run(para, text, bold=False, italic=False, size=Pt(10), color=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = size
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def add_inline(para, text, base_size=Pt(10), base_bold=False):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            styled_run(para, part[2:-2], bold=True, size=base_size)
        else:
            styled_run(para, part, bold=base_bold, size=base_size)

with open(MD_FILE, encoding="utf-8") as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    raw = lines[i].rstrip("\n")
    stripped = raw.strip()

    if stripped == "" or stripped.startswith("<div") or stripped == "</div>":
        i += 1
        continue

    # Section Break for 2-column clauses
    if stripped.startswith("**1. Client obligations:"):
        new_sect = doc.add_section(WD_SECTION.CONTINUOUS)
        setup_section(new_sect)
        set_columns(new_sect, 2)
        
    # Section Break back to 1-column for Signatory Page
    if stripped.startswith("IN WITNESS WHEREOF"):
        new_sect = doc.add_section(WD_SECTION.CONTINUOUS)
        setup_section(new_sect)
        set_columns(new_sect, 1)

    # Logo injection
    if stripped.startswith("# ![Softway Logo]"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(24)
        if os.path.exists(LOGO_FILE):
            r = p.add_run()
            r.add_picture(LOGO_FILE, width=Inches(3.0))
        else:
            run = p.add_run("[ LOGO.PNG MISSING ]")
            run.bold = True
        i += 1
        continue

    # Main Title
    if stripped == "MASTER SERVICES AGREEMENT":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(18)
        run = p.add_run(stripped)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(14)
        i += 1
        continue

    # Signatory Table
    if stripped.startswith("|"):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_lines.append(lines[i].strip())
            i += 1
        rows = [r for r in table_lines if not re.match(r"^\|[-| :]+\|$", r)]
        if rows:
            parsed = []
            for row in rows:
                cols = [c.strip() for c in row.strip("|").split("|")]
                parsed.append(cols)
            ncols = max(len(r) for r in parsed)
            tbl = doc.add_table(rows=len(parsed), cols=ncols)
            tbl.style = "Table Grid"
            for ri, row in enumerate(parsed):
                for ci, cell_text in enumerate(row):
                    cell = tbl.cell(ri, ci)
                    cell.paragraphs[0].clear()
                    add_inline(cell.paragraphs[0], cell_text, base_size=Pt(10))
                    cell.paragraphs[0].paragraph_format.space_after = Pt(4)
            doc.add_paragraph()
        continue

    # Regular text paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

    # Handling hanging indents (0.5 inch hanging indent for sub-clauses)
    is_hanging = False
    if re.match(r"^\s*[A-Z]\.", raw) or re.match(r"^\s*\*\*\d+\.\d+", raw) or (i > 0 and lines[i-1].strip() == "" and raw.startswith("   ")):
        is_hanging = True
        p.paragraph_format.left_indent = Inches(0.5)
        # Apply the -0.5 inch first line indent so it 'hangs' to the left if the text wraps
        p.paragraph_format.first_line_indent = Inches(-0.5)

    # Clause Headings
    full_bold_match = re.match(r"^\*\*(.+)\*\*$", stripped)
    if full_bold_match:
        run = p.add_run(full_bold_match.group(1))
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(10)
    else:
        add_inline(p, stripped)

    i += 1

doc.save(OUT_FILE)
print("DOCX generated successfully with hanging indents!")
